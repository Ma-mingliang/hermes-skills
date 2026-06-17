"""
Verify a filled .docx document against requirements.
Usage: python verify_docx.py <docx_path> [--template <template_path>]

Checks:
  1. Structure (rows/columns match template)
  2. Images are embedded (not just caption text)
  3. Chinese font settings (eastAsia = 宋体)
  4. Line spacing (1.5x)
  5. Content completeness (keyword check)
"""

import sys
import os
from docx import Document
from docx.oxml.ns import qn


def check_images(cell):
    """Count embedded images and check captions have images."""
    img_count = 0
    missing = []
    for pi, p in enumerate(cell.paragraphs):
        blips = p._element.findall('.//' + qn('a:blip'))
        if blips:
            img_count += 1
        text = p.text.strip()
        if text and ('图' in text[:5] or '图 ' in text[:5]):
            prev_has_img = False
            if pi > 0:
                prev_has_img = bool(
                    cell.paragraphs[pi - 1]._element.findall('.//' + qn('a:blip'))
                )
            if not prev_has_img:
                missing.append(f"[{pi}] {text[:60]}")
    return img_count, missing


def check_font(cell):
    """Check that eastAsia font is 宋体."""
    issues = []
    for pi, p in enumerate(cell.paragraphs):
        for run in p.runs:
            rpr = run._element.find(qn('w:rPr'))
            if rpr is not None:
                for f in rpr.findall(qn('w:rFonts')):
                    ea = f.get(qn('w:eastAsia'))
                    if ea and ea != '宋体':
                        issues.append(f"[{pi}] eastAsia={ea}")
    return issues


def check_line_spacing(cell):
    """Check that paragraphs have 1.5x line spacing."""
    issues = []
    for pi, p in enumerate(cell.paragraphs):
        if p.text.strip() and p.paragraph_format.line_spacing != 1.5:
            issues.append(
                f"[{pi}] spacing={p.paragraph_format.line_spacing}"
            )
    return issues


def verify(docx_path, template_path=None, keywords=None):
    doc = Document(docx_path)
    table = doc.tables[0]

    print(f"=== 验证: {os.path.basename(docx_path)} ===")
    print(f"表格: {len(table.rows)}行 x {len(table.columns)}列")

    # Structure check against template
    if template_path and os.path.exists(template_path):
        tmpl = Document(template_path)
        tt = tmpl.tables[0]
        if len(table.rows) != len(tt.rows):
            print(f"  ✗ 行数不匹配: {len(table.rows)} vs template {len(tt.rows)}")
        else:
            print(f"  ✓ 行数匹配: {len(table.rows)}")

    total_imgs = 0
    total_missing = []
    total_font_issues = []
    total_spacing_issues = []

    for ri, row in enumerate(table.rows):
        cell = row.cells[0]
        imgs, missing = check_images(cell)
        font_issues = check_font(cell)
        spacing_issues = check_line_spacing(cell)

        total_imgs += imgs
        total_missing.extend([(ri, m) for m in missing])
        total_font_issues.extend([(ri, f) for f in font_issues])
        total_spacing_issues.extend([(ri, s) for s in spacing_issues])

        print(f"  行{ri}: {len(cell.paragraphs)}段, {imgs}张图片")

    # Report
    print(f"\n图片: {total_imgs}张嵌入")
    for ri, m in total_missing:
        print(f"  ✗ 行{ri} 缺图: {m}")

    print(f"\n字体检查:")
    if total_font_issues:
        for ri, f in total_font_issues:
            print(f"  ✗ 行{ri} {f}")
    else:
        print("  ✓ 全部宋体")

    print(f"\n行距检查:")
    if total_spacing_issues:
        for ri, s in total_spacing_issues[:5]:
            print(f"  ✗ 行{ri} {s}")
    else:
        print("  ✓ 全部1.5倍行距")

    # Keyword check
    if keywords:
        print(f"\n内容关键词检查:")
        full_text = ' '.join(p.text for cell in [r.cells[0] for r in table.rows]
                             for p in cell.paragraphs)
        for kw in keywords:
            found = kw in full_text
            print(f"  {'✓' if found else '✗'} {kw}")

    # File size
    size = os.path.getsize(docx_path)
    print(f"\n文件大小: {size / 1024:.0f} KB")
    print("=== 验证完成 ===")


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('docx_path')
    parser.add_argument('--template', default=None)
    parser.add_argument('--keywords', nargs='*', default=None)
    args = parser.parse_args()
    verify(args.docx_path, args.template, args.keywords)
